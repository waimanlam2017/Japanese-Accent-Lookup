from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from collections import deque
import MeCab
import re
import sys 
import urllib.request
import urllib.parse
import pickle
from urllib.parse   import quote

WEBLIO_URL = "http://www.weblio.jp/content/"

class JapanTextAnalyzer():

	def __init__(self):
		self.japan_pos = ['連体詞', '接続詞', '助詞', '形容詞', '記号', '名詞', '接頭詞', '副詞', 'BOS/EOS', '助動詞', 'フィラー', '感動詞', '動詞']
		self.target_post = ['形容詞', '名詞']
		self.punc_post = ['記号']
		self.aiueo = ['あ','い','う','え','お']
		self.cache_dict = {}
		self.doc_filename = 'processed_japan_text.doc'
		self.source_filename = 'japan_text.txt'
		
		
	def parse_text(self, text):
		m = MeCab.Tagger("-Ochasen")		
		m.parse("")
		parsed = m.parseToNode(text)
		components = []
		while parsed:
			word = parsed.surface
			pos = parsed.feature.split(',')[0]
			tp = (word, pos)		
			components.append(tp)
			parsed = parsed.next
		return components

	def main(self):
		# Source text file - Start 
		f = open('japan_text.txt', encoding='utf-8')
		text = f.read()
		f.close()
		# Source text file - End
		
		# MeCab Tagging Process - Start
		tagged_text_tp = self.parse_text(text)
		# MeCab Tagging Process - End
		
		# Extract Accent - Start
		output_line = ''
		debug_note = []
		no_result_note = []
		accent_note = []
		for word in tagged_text_tp:			
			original_word = word[0]
			pos = word[1]
			result = None
			if ( pos in self.target_post ): #Lookup dictionary
				if ( pos == "形容詞" and ( original_word.endswith('かっ') or original_word.endswith('く') ) ):
					result = self.weblio_special_search(original_word, pos)
				else:
					result = self.weblio(original_word, pos)
				if ( result ):#Dictionary returned something
					accent = result[0]
					debug_info = result[2]
					
					if ( "イ形容詞變化型" in debug_info ):
						accent_note.append(result[2])
					elif ( debug_info != '' ):
						debug_note.append(result[2])
											
					if ( accent == 0 ):
						symbol_accent = 'α'
					elif ( accent == 1 ):
						symbol_accent = 'β'
					elif ( accent == 2 ):
						symbol_accent = 'γ'
					elif ( accent == 3 ):
						symbol_accent = 'δ'
					elif ( accent == 4 ):
						symbol_accent = 'ϵ'
					elif ( accent == 5 ):
						symbol_accent = 'ζ'
					elif ( accent == 6 ):
						symbol_accent = 'η'
					else:
						symbol_accent = accent
					output_line += str(symbol_accent) + original_word				
				else:#Dictionary no result
					no_result_note.append(original_word + " : " + pos + ", 字典查無此字。")
					output_line += original_word				
			else:# Does not Lookup Dictionary				
				output_line += original_word
				
			if ( pos in self.punc_post and ( original_word == '。' or original_word == '？' or original_word == '?' ) ):
				output_line += "Ω"
			
		#Extract Accent - End		
		return (output_line, accent_note, debug_note, no_result_note)
		
	def write_to_doc(self, output_line, accent_note, debug_note, no_result_note):
		accent_symbol = ['α','β','γ','δ','ϵ','ζ','η']
		first_no_voice = ['き','し','ち','ひ','ぴ','く','す','つ','ふ','ぷ']
		second_no_voice = ['か','き','く','け','こ','さ','し','す','せ','そ','た','ち','つ','て','と','は','ひ','ふ','へ','ほ','ぱ','ぴ','ぷ','ぺ','ぽ']
		run_stack = []
		document = Document()
		document.add_heading('Japan')
		paragraph = document.add_paragraph()
		for char in output_line:
			if ( char in accent_symbol ):
				if ( char == 'α' ):
					accent = '0'
				elif ( char == 'β' ):
					accent = '1'
				elif ( char == 'γ' ):
					accent = '2'
				elif ( char == 'δ' ):
					accent = '3'
				elif ( char == 'ϵ' ):
					accent = '4'
				elif ( char == 'ζ' ):
					accent = '5'
				elif ( char == 'η' ):
					accent = '6'
				run = paragraph.add_run(accent)
				run.font.superscript = True			
			elif ( char == 'Ω' ):
				paragraph = document.add_paragraph()
				continue
			else:
				run = paragraph.add_run(char)
			run_stack.append(run)
			
		run_stack = deque(run_stack)
			
		while len(run_stack) > 1:
			if ( run_stack[1]._r.text == 'っ' ):
				if ( run_stack[0]._r.text in first_no_voice and run_stack[2]._r.text in second_no_voice ):
					run_stack[0].font.highlight_color = WD_COLOR_INDEX.GRAY_25
				run_stack.popleft() 
				run_stack.popleft() 
			elif ( run_stack[0]._r.text in first_no_voice and run_stack[1]._r.text in second_no_voice ):
				run_stack[0].font.highlight_color = WD_COLOR_INDEX.GRAY_25
				run_stack.popleft() 
			else:
				run_stack.popleft() 
		

		document.add_page_break()
		document.add_heading('Dictionary Note')
		for line in debug_note:
			paragraph = document.add_paragraph()
			run = paragraph.add_run(line)		
			if ( "請覆查字典" in line ):
				run.font.bold = True
				
		document.add_page_break()
		document.add_heading('Accent Note')
		for line in accent_note:
			paragraph = document.add_paragraph()
			paragraph.add_run(line)
			
		document.add_page_break()
		document.add_heading('No Result Note')
		for line in no_result_note:
			paragraph = document.add_paragraph()
			paragraph.add_run(line)
		
		document.save(self.doc_filename)	
					
		
	def weblio_special_search(self, word, pos):		
		print("Special checking with " + word )
		#Checking variation with く - Second Pass - Start
		if ( pos == "形容詞" and word.endswith("く") ):
			'''原形形容詞的重音如果是
			◎、則重音依然是◎、若不
			是◎、則在「く」的前面兩
			個字。'''
			transformed_word = word
			none_changing_part = transformed_word[:-1]
			original_word = none_changing_part + "い"
			
			#Query Weblio with original Adjective
			result = self.weblio(original_word, pos)
			#Query Weblio with original Adjective
			
			if ( result ):
				original_accent = int(result[0])
				pronunciation = result[1]
				pronunciation_non_changing_part = pronunciation[:-1]
				
				if ( original_accent == 0 ):#Special rule
					vary_accent = 0							
				else:#Special rule
					vary_accent = original_accent-1 if original_accent > 1 else 1
				
				debug_line = "イ形容詞變化型-く: " + transformed_word + ", 發音: " + pronunciation_non_changing_part + "く, 聲調(按規則推斷): " + str(vary_accent) + ", 原來聲調: " + str(original_accent)
				
				self.cache_dict[word] = [vary_accent, pronunciation, debug_line, word]
				return (vary_accent, pronunciation, debug_line, word)
			else:
				return None
		elif ( pos == "形容詞" and word.endswith("かっ") ): #Checking variation with かっ 
			'''原形形容詞的重音如果是
			◎、則重音在「かった」的
			前面一個字、若不是◎、則
			在「かった」的前面兩個字。'''			
			transformed_word = word
			none_changing_part = transformed_word[:-2]
			original_word = none_changing_part + "い"
			
			#Query Weblio with original Adjective
			result = self.weblio(original_word, pos)
			#Query Weblio with original Adjective
			
			if ( result ):
				original_accent = result[0]
				pronunciation = result[1]
				pronunciation_non_changing_part = pronunciation[:-1]
				
				if ( original_accent == 0 ):#Special rule
					if ( pronunciation_non_changing_part[-1] in self.aiueo ):#Accent is あいうえお Special rule
						vary_accent = str(len(pronunciation_non_changing_part)-2)
					else:#Special rule
						vary_accent = str(len(pronunciation_non_changing_part)-1)
				else:#Special rule
					vary_accent = original_accent-1 if original_accent > 1 else 1
				
				debug_line = "イ形容詞變化型-かった: " + transformed_word + ", 發音: " + pronunciation_non_changing_part + "かった, 聲調(估計): " + str(vary_accent) + ", 原來聲調: " + str(original_accent)
				
				self.cache_dict[word] = [vary_accent, pronunciation, debug_line, word]
				return ( vary_accent, pronunciation, debug_line, word )
			else:
				return None
		else:#Other transformed adjective, not supported now
			return None
		
	
	def weblio(self, word, pos):
		result = self.special_word_checking(word, pos)
		if ( result is not None ):
			return ( result[0], result[1], result[2], result[3] )
	
		if ( word in self.cache_dict ):
			result = self.cache_dict[word]
			return ( result[0], result[1], result[2], result[3] )
	
		print("Normal Looking up dictionary for " + word )
		filename, headers = urllib.request.urlretrieve(WEBLIO_URL + quote(word))
		with open(filename, encoding='utf-8') as f:
			html_doc = f.read()
			soup = BeautifulSoup(html_doc, 'html.parser')
			
			#Checking the pronunciation - Start
			pronunciation_found = False
			pronunciation = ''
			div_results = soup.find_all("div", "NetDicHead")
			for result in div_results:
				if ( pronunciation_found ):
					break
				b_results = result.find_all("b")
				for result in b_results:
					search_result = re.findall(r'\D+', result.string)
					if ( len(search_result) > 0 ):
						pronunciation_found = True
						pronunciation = ''.join(search_result)			
						break
			#Checking the pronunciation - End
			
			#Checking the accent - First Pass Start
			accent_found = False
			accent = ''
			debug_line = ''
			accent_result_count = 0
			div_results = soup.find_all("div", "NetDicHead")
			for result in div_results:
				if ( accent_found ):
					break
				span_results = result.find_all("span")
				for result in span_results:
					search_result = re.findall(r'\d+', result.string)
					if ( len(search_result) > 0 ):
						accent_found = True
						accent_result_count += 1
						accent = int(''.join(search_result))
						debug_line = word + ": " + pos + ", 發音 : " + pronunciation + ", 聲調: " + str(accent)
						break
						
			if ( accent_result_count > 1 ):
				debug_line += ". 請覆查字典。檢索結果多於1。共有" + str(accent_result_count) + "個檢索結果。" 
				
				
			if ( accent_found ):
				self.cache_dict[word] = [accent, pronunciation, debug_line, word]
				return ( accent, pronunciation, debug_line, word )			
			else:
				return None
				
	def special_word_checking(self, word, pos):
		#if ( 'よく' in word and '形容詞' in pos ):
		#	return ( '1', 'よく', word + ": " + pos + ', 修正的特例 - よい的過去型', 'よく' )
		#else:
		#	return None
		pass
	
	def save_cache_dict(self, pickle_name):
		with open(pickle_name, 'wb') as handle:
			pickle.dump(self.cache_dict, handle, protocol=pickle.HIGHEST_PROTOCOL)
				
	def load_cache_dict(self, pickle_name):
		with open(pickle_name, 'rb') as handle:
			self.cache_dict = pickle.load(handle)
	
if __name__ == "__main__":
	Jp = JapanTextAnalyzer()
	#Jp.load_cache_dict('pickle_dict.dict')
	result = Jp.main()
	Jp.write_to_doc(result[0], result[1], result[2], result[3])
	Jp.save_cache_dict('pickle_dict.dict')

	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	
	